from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Laporan_Praktikum_Cloudify.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.1
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9E2EC")


def set_cell_width(cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def apply_base_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Laporan Praktikum Cloudify")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(102, 112, 133)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("LAPORAN PRAKTIKUM")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("Mini Project Website Cloud Storage - Cloudify")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(46, 116, 181)

    meta = doc.add_table(rows=5, cols=2)
    set_table_borders(meta)
    widths = [1.75, 4.75]
    rows = [
        ("Nama", "........................................................"),
        ("NIM", "........................................................"),
        ("Kelas", "........................................................"),
        ("Mata Praktikum", "Mini Project Website"),
        ("Tanggal", "24 Mei 2026"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
        set_cell_text(row.cells[0], label, bold=True)
        set_cell_text(row.cells[1], value)
    doc.add_paragraph()


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    set_table_borders(table)
    table.rows[0].cells[0].width = Inches(1.9)
    table.rows[0].cells[1].width = Inches(4.6)
    for i, (label, value) in enumerate(rows):
        if i > 0:
            table.add_row()
        cells = table.rows[i].cells
        set_cell_width(cells[0], 1.9)
        set_cell_width(cells[1], 4.6)
        set_cell_text(cells[0], label, bold=True)
        set_cell_text(cells[1], value)
    doc.add_paragraph()


def add_matrix_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "F2F4F7")
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            set_cell_width(cell, widths[idx])
            set_cell_text(cell, value)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def build() -> None:
    doc = Document()
    apply_base_styles(doc)
    add_footer(doc)
    add_title(doc)

    doc.add_heading("1. Pendahuluan", level=1)
    doc.add_paragraph(
        "Cloudify adalah website cloud storage lokal berbasis PHP yang digunakan untuk "
        "mengelola file gambar melalui antarmuka web. Website ini dibuat sebagai mini "
        "project praktikum dengan fokus pada halaman utama, manajemen pengguna berbasis "
        "role, pembatasan akses file, serta validasi upload berdasarkan ukuran, kapasitas, "
        "ekstensi, dan MIME type."
    )
    doc.add_paragraph(
        "Tujuan praktikum ini adalah menerapkan konsep aplikasi web dinamis yang mampu "
        "menerima unggahan file, menampilkan file dalam bentuk galeri, mengatur hak akses "
        "pengguna, dan menjaga keamanan dasar melalui autentikasi, CSRF token, sanitasi "
        "nama file, serta pencatatan audit."
    )

    doc.add_heading("2. Identitas Sistem", level=1)
    add_key_value_table(
        doc,
        [
            ("Nama aplikasi", "Cloudify"),
            ("Jenis aplikasi", "Website cloud storage / galeri file lokal"),
            ("Teknologi", "PHP, MySQL/MariaDB, HTML, CSS, JavaScript, XAMPP"),
            ("Folder upload", "uploads/"),
            ("Database", "simple_cloud"),
            ("Zona waktu", "Asia/Makassar"),
        ],
    )

    doc.add_heading("3. Pemetaan Kebutuhan Tugas", level=1)
    add_matrix_table(
        doc,
        ["No", "Kebutuhan Tugas", "Implementasi Pada Website", "Status"],
        [
            [
                "1",
                "Membuat home page website",
                "Tersedia pada index.php berupa landing/gallery page publik dengan preview gambar dan akses menuju login atau dashboard.",
                "Terpenuhi",
            ],
            [
                "2",
                "Admin mengelola semua pengguna dan semua file",
                "Role superadmin/admin dapat mengakses seluruh file. Superadmin memiliki modul manage_users melalui manage_user.php.",
                "Terpenuhi, dengan catatan admin biasa fokus pada semua file sedangkan superadmin mengelola user.",
            ],
            [
                "3",
                "User upload, download, dan hapus file miliknya sendiri",
                "Role user memiliki permission upload, download, delete, rename. FileOwnershipStore membatasi akses berdasarkan owner_id.",
                "Terpenuhi",
            ],
            [
                "4",
                "Viewer hanya melihat dan mengunduh file tertentu",
                "Role viewer belum dibuat eksplisit. Fungsi serupa dapat dipenuhi oleh public gallery/download public atau perlu penambahan role viewer pada AuthManager, UserStore, dan database enum.",
                "Perlu penyelarasan",
            ],
            [
                "5",
                "Guest akses terbatas atau hanya melihat halaman tertentu",
                "Role guest hanya memiliki permission home dan dashboard. Guest tidak dapat upload, download internal, rename, atau delete file.",
                "Terpenuhi",
            ],
            [
                "6",
                "Batas kapasitas dan jenis file upload",
                "config/app.php membatasi ukuran file 10 MB, kapasitas storage 1 GB, dan whitelist png, jpg, jpeg, gif, webp disertai validasi MIME.",
                "Terpenuhi",
            ],
        ],
        [0.45, 1.65, 3.25, 1.15],
    )

    doc.add_heading("4. Analisis Fitur", level=1)
    doc.add_heading("4.1 Home Page", level=2)
    doc.add_paragraph(
        "Halaman utama pada index.php berfungsi sebagai etalase publik Cloudify. "
        "Konten file ditampilkan dalam bentuk galeri gambar sehingga pengguna dapat "
        "melihat isi storage tanpa langsung masuk ke dashboard. Halaman ini juga "
        "menyediakan preview modal dan tombol download publik untuk file gambar yang "
        "ditampilkan."
    )
    add_bullets(
        doc,
        [
            "Menjadi halaman pertama yang memperkenalkan aplikasi dan isi storage.",
            "Menampilkan file gambar dari folder uploads dengan metadata ukuran dan waktu modifikasi.",
            "Mengarahkan pengguna login ke dashboard dan pengguna baru ke halaman autentikasi.",
        ],
    )

    doc.add_heading("4.2 Autentikasi dan Level Pengguna", level=2)
    doc.add_paragraph(
        "Pengaturan hak akses utama berada pada AuthManager::roleAllows(). Sistem "
        "menggunakan session untuk menyimpan id, nama, dan role pengguna setelah login. "
        "Password diverifikasi dengan password_verify sehingga password tidak dibandingkan "
        "dalam bentuk teks biasa."
    )
    add_matrix_table(
        doc,
        ["Role", "Hak Akses Utama", "Penjelasan"],
        [
            ["Superadmin", "Semua modul termasuk manage_users", "Mengelola struktur pengguna, file, audit, upload, download, rename, dan delete."],
            ["Admin", "Semua file dan audit", "Mengelola semua file di storage, termasuk rename, download, dan delete."],
            ["User", "File milik sendiri", "Dapat upload dan mengelola file dengan owner_id yang sama dengan akunnya."],
            ["Guest", "Home dan dashboard terbatas", "Tidak memiliki akses upload, download internal, rename, atau delete."],
            ["Viewer", "Belum eksplisit", "Sebaiknya ditambahkan untuk memenuhi instruksi tugas secara penuh."],
        ],
        [1.05, 2.0, 3.45],
    )

    doc.add_heading("4.3 Upload dan Validasi File", level=2)
    doc.add_paragraph(
        "Proses upload dilakukan melalui upload.php dan diproses oleh CloudStorageService. "
        "Validasi dilakukan berlapis agar file yang masuk sesuai kebijakan sistem."
    )
    add_numbered(
        doc,
        [
            "Memastikan request upload menggunakan metode POST.",
            "Memvalidasi CSRF token agar aksi upload berasal dari form yang sah.",
            "Memeriksa error upload dari PHP, ukuran file, kapasitas storage, dan sumber temporary file.",
            "Membersihkan nama file agar tidak mengandung karakter berbahaya atau path traversal.",
            "Memeriksa ekstensi file terhadap whitelist png, jpg, jpeg, gif, dan webp.",
            "Memeriksa MIME type agar isi file sesuai dengan ekstensi yang diklaim.",
            "Mencatat pemilik file ke tabel files dan menulis event upload ke audit log.",
        ],
    )
    add_key_value_table(
        doc,
        [
            ("Maksimum per file", "10 MB"),
            ("Maksimum storage", "1 GB"),
            ("Ekstensi diizinkan", "png, jpg, jpeg, gif, webp"),
            ("Validasi tambahan", "MIME type, ukuran, kapasitas, nama file aman, dan file upload valid"),
        ],
    )

    doc.add_heading("4.4 Manajemen File", level=2)
    doc.add_paragraph(
        "Dashboard menyediakan daftar file dalam mode list dan grid. Pengguna dapat "
        "melakukan download, rename, dan delete sesuai hak akses. Admin/superadmin dapat "
        "melihat semua file, sedangkan user hanya melihat file yang tercatat sebagai miliknya."
    )
    add_bullets(
        doc,
        [
            "Download melalui download.php dengan pengecekan permission dan kepemilikan file.",
            "Delete melalui delete.php hanya menerima POST dan wajib CSRF token.",
            "Rename menjaga ekstensi file agar tidak berubah dan nama tujuan tidak bentrok.",
            "Preview gambar disediakan melalui preview.php untuk MIME gambar yang valid.",
        ],
    )

    doc.add_heading("4.5 Audit dan Keamanan", level=2)
    doc.add_paragraph(
        "Website mencatat aktivitas penting seperti upload, download, delete, dan manajemen "
        "pengguna. AuditLogger menyimpan timestamp, aksi, status, user, IP, user agent, dan "
        "konteks tambahan ke logs/audit.log. Dari sisi keamanan, sistem memakai CSRF token, "
        "session login, sanitasi nama file, validasi MIME type, dan pembatasan akses berbasis role."
    )

    doc.add_heading("5. Alur Kerja Sistem", level=1)
    add_numbered(
        doc,
        [
            "Pengunjung membuka home page Cloudify dan melihat galeri publik.",
            "Pengguna login melalui login.php menggunakan akun yang tersimpan di database.",
            "AuthManager menyimpan session dan mengecek permission role.",
            "Pengguna yang memiliki izin upload memilih file gambar dari dashboard.",
            "CloudStorageService memvalidasi ukuran, kapasitas, ekstensi, MIME, dan nama file.",
            "File yang valid dipindahkan ke uploads/ dan metadata owner dicatat ke tabel files.",
            "Aksi pengguna seperti upload, download, rename, dan delete dicatat ke audit log.",
        ],
    )

    doc.add_heading("6. Pengujian Fitur", level=1)
    add_matrix_table(
        doc,
        ["Skenario", "Langkah Uji", "Hasil yang Diharapkan"],
        [
            ["Home page", "Buka http://localhost/simple_cloud/", "Galeri publik tampil dan tombol login/dashboard tersedia."],
            ["Login admin", "Masuk sebagai superadmin/admin.", "Dashboard menampilkan seluruh file dan kontrol manajemen sesuai role."],
            ["Login user", "Masuk sebagai user lalu upload gambar valid.", "File berhasil masuk dan owner_id user tercatat."],
            ["Guest", "Masuk sebagai guest.", "Panel upload tidak tersedia dan akses file dibatasi."],
            ["Upload melebihi batas", "Upload file lebih dari 10 MB.", "Sistem menolak dengan pesan ukuran melebihi batas."],
            ["Upload ekstensi salah", "Upload file selain png, jpg, jpeg, gif, webp.", "Sistem menolak karena ekstensi/MIME tidak diizinkan."],
            ["Delete file", "User mencoba hapus file miliknya.", "File terhapus, metadata dihapus, dan audit log tercatat."],
        ],
        [1.25, 2.65, 2.6],
    )

    doc.add_heading("7. Evaluasi dan Rekomendasi", level=1)
    doc.add_paragraph(
        "Secara umum Cloudify sudah memenuhi sebagian besar batasan mini project. "
        "Fitur home page, upload file, pembatasan ukuran dan tipe file, dashboard, audit, "
        "serta pembatasan akses admin, user, dan guest telah tersedia. Catatan utama "
        "berada pada role viewer karena instruksi tugas menyebut viewer sebagai level "
        "pengguna tersendiri, sedangkan kode saat ini belum mendefinisikan role tersebut."
    )
    add_bullets(
        doc,
        [
            "Tambahkan role viewer ke enum database users, UserStore::ROLES, dan AuthManager::roleAllows().",
            "Berikan viewer permission dashboard dan download tanpa upload, rename, atau delete.",
            "Tambahkan mekanisme daftar file tertentu untuk viewer, misalnya tabel file_shares berisi viewer_id dan file_name.",
            "Sesuaikan halaman manajemen user agar superadmin dapat memilih role viewer.",
            "Update README agar batas upload mengikuti konfigurasi terbaru, yaitu 10 MB dan hanya file gambar.",
        ],
    )

    doc.add_heading("8. Kesimpulan", level=1)
    doc.add_paragraph(
        "Mini project Cloudify menunjukkan penerapan website cloud storage sederhana "
        "dengan antarmuka home page, dashboard manajemen file, autentikasi, role-based "
        "access control, validasi upload, dan audit aktivitas. Sistem sudah layak sebagai "
        "praktikum dasar manajemen file berbasis web. Agar sepenuhnya selaras dengan "
        "instruksi tugas, pengembangan berikutnya perlu menambahkan role viewer eksplisit "
        "yang hanya dapat melihat dan mengunduh file tertentu."
    )

    doc.add_heading("Lampiran: File Penting", level=1)
    add_matrix_table(
        doc,
        ["File", "Fungsi"],
        [
            ["index.php", "Home page publik dan galeri visual."],
            ["dashboard.php", "Dashboard utama setelah login."],
            ["upload.php", "Endpoint upload file."],
            ["download.php", "Endpoint download dan inline preview."],
            ["delete.php", "Endpoint hapus file dengan POST dan CSRF."],
            ["src/Security/AuthManager.php", "Autentikasi dan permission role."],
            ["src/Services/CloudStorageService.php", "Storage, validasi upload, list, rename, delete."],
            ["src/Services/FileOwnershipStore.php", "Pengecekan kepemilikan dan metadata file."],
            ["src/Services/UserStore.php", "Manajemen akun pengguna dari database."],
            ["config/app.php", "Konfigurasi batas file, kapasitas, dan MIME whitelist."],
        ],
        [2.35, 4.15],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
